#!/usr/bin/env python3
"""
bolt_cert.py - Generate Dunsteel bolt tightening documents from PM16 + PM17 templates.

For a given project number and scope, this:
  1. Locates the project folder on the S: drive ("[number] -*").
  2. Resolves the QA scope folder ([project]\\08 QA\\[scope]).
  3. Copies the PM17 (Bolt Tightening Certificate) and PM16 (Bolt Tightening
     Methods) templates into that folder, renamed per the Dunsteel convention.
  4. Fills placeholders via python-docx (head contractor, project name, scope,
     signatory, role, phone, date).
  5. Renders a PDF for each via docx2pdf (Word COM, native quality).

Two separate documents are always produced (4 files total):
  PM17  ->  "[number] - Bolt tightening Certificate - [scope].docx" (+ .pdf)
  PM16  ->  "[number] - Bolt Tightening Methods - [scope].docx" (+ .pdf)

This is the worker called by the /bolt-cert skill. It does NOT prompt; the skill
layer handles interactive confirmation (missing scope folder, ambiguous scope).
The script's job is the safe, deterministic file work.

Usage:
    python scripts/bolt_cert.py --project 501 --scope "Level 2"
    python scripts/bolt_cert.py --project 501 --scope "Level 2" --dest-subfolder TEST_boltcert
    python scripts/bolt_cert.py --project 501 --list-scopes
    python scripts/bolt_cert.py --project 501 --scope "Stair 1" \\
        --contractor "Northbridge Constructions" --project-name "Riverside Eastern" \\
        --signatory "Nathan Hancock" --role "Project Manager" --phone "04XX XXX XXX"

Notes:
  - HARD RULE: no long dashes anywhere in code or generated output. Hyphens only.
  - --dest-subfolder writes into [scope]\\<sub> instead of [scope] itself
    (used by the test harness to avoid touching live cert folders).
  - --create-scope creates the scope folder if missing (the skill only passes
    this after the user confirms in interactive use).
"""

import argparse
import datetime as _dt
import shutil
import sys
from pathlib import Path

from docx import Document

# --- Fixed paths -------------------------------------------------------------

CURRENT_PROJECTS_ROOT = Path(r"S:\Operations\01 Current Project")
TEMPLATES_ROOT = Path(
    r"S:\Operations\11 Project Management Database\Project Management Templates"
)
PM17_TEMPLATE = TEMPLATES_ROOT / "PM17. Bolt tightening Certificate.docx"
PM16_TEMPLATE = TEMPLATES_ROOT / "PM16. Bolt Tighteneing Methods.docx"

# Default signatory (overridable via CLI).
DEFAULT_SIGNATORY = "Nathan Hancock"
DEFAULT_ROLE = "Project Manager"
DEFAULT_PHONE = "04XX XXX XXX"

# The literal placeholder token used in both templates.
PLACEHOLDER = "???"


# --- Project / scope resolution ---------------------------------------------

def find_project_folder(project_number: str) -> Path:
    """Return the S: drive folder whose name starts with '[number] -'."""
    if not CURRENT_PROJECTS_ROOT.exists():
        raise FileNotFoundError(
            f"Current Projects root not found: {CURRENT_PROJECTS_ROOT}"
        )
    matches = [
        p for p in CURRENT_PROJECTS_ROOT.iterdir()
        if p.is_dir() and p.name.startswith(f"{project_number} -")
    ]
    if not matches:
        raise FileNotFoundError(
            f"No project folder matching '{project_number} -*' under "
            f"{CURRENT_PROJECTS_ROOT}"
        )
    if len(matches) > 1:
        names = ", ".join(m.name for m in matches)
        raise ValueError(
            f"Multiple project folders match '{project_number} -*': {names}. "
            "Disambiguate before running."
        )
    return matches[0]


def qa_root(project_folder: Path) -> Path:
    return project_folder / "08 QA"


def list_scope_folders(project_folder: Path) -> list[str]:
    """List sub-folders inside 08 QA (the available scopes)."""
    qa = qa_root(project_folder)
    if not qa.exists():
        return []
    return sorted(p.name for p in qa.iterdir() if p.is_dir())


def resolve_scope_folder(project_folder: Path, scope: str,
                         create: bool) -> Path:
    """Resolve [project]\\08 QA\\[scope], creating it only if create=True."""
    qa = qa_root(project_folder)
    if not qa.exists():
        raise FileNotFoundError(f"QA folder not found: {qa}")
    scope_folder = qa / scope
    if not scope_folder.exists():
        if not create:
            raise FileNotFoundError(
                f"Scope folder not found: {scope_folder}. "
                "Pass --create-scope to create it (after user confirmation)."
            )
        scope_folder.mkdir(parents=True, exist_ok=False)
    return scope_folder


# --- Placeholder filling -----------------------------------------------------

def _replace_placeholder_in_paragraph(paragraph, value: str) -> bool:
    """
    Replace the '???' token within a paragraph, preserving formatting.

    The token sits in a single run in both templates, so a per-run replace
    works. As a fallback (token split across runs) we collapse runs.
    """
    replaced = False
    for run in paragraph.runs:
        if PLACEHOLDER in run.text:
            run.text = run.text.replace(PLACEHOLDER, value)
            replaced = True
    if not replaced and PLACEHOLDER in paragraph.text:
        # Fallback: rebuild from full text into the first run.
        new_text = paragraph.text.replace(PLACEHOLDER, value)
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""
        replaced = True
    return replaced


def _set_paragraph_text_preserve_format(paragraph, value: str) -> None:
    """Overwrite a paragraph's text, keeping the first run's formatting."""
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)


def fill_certificate(template_path: Path, out_path: Path, *,
                     contractor: str, project_name: str, scope: str,
                     signatory: str, role: str, phone: str,
                     doc_date: str) -> None:
    """
    Copy + fill the PM17 Bolt Tightening Certificate.

    Placeholder map (paragraph indices confirmed against the live template):
      P13  'Attn: ???'                              -> contractor
      P15  '... at the ??? project.'                -> project name
      P32  signatory name (template default name)   -> signatory
      P33  role line (empty in template)            -> role
      P34  phone (template default number)          -> phone

    The project name carries the scope reference (e.g. "Riverside Eastern -
    Level 2") so the certificate is scope-specific, matching the existing
    501 cert convention. A dated line is added under the title.
    """
    shutil.copy2(template_path, out_path)
    doc = Document(str(out_path))
    paras = doc.paragraphs

    # Project name on the cert includes the scope reference.
    project_ref = f"{project_name} - {scope}" if scope else project_name

    # P13: Attn line -> head contractor.
    _replace_placeholder_in_paragraph(paras[13], contractor)
    # P15: regarding line -> project + scope.
    _replace_placeholder_in_paragraph(paras[15], project_ref)

    # Signatory block. Indices are stable in the template; guard anyway.
    if len(paras) > 32:
        _set_paragraph_text_preserve_format(paras[32], signatory)
    if len(paras) > 33:
        # Role line is blank in the template; set it (matching the 501 cert).
        if paras[33].runs:
            _set_paragraph_text_preserve_format(paras[33], role)
        else:
            r = paras[33].add_run(role)
            # Inherit basic look from the name run if present.
            if paras[32].runs:
                src = paras[32].runs[0].font
                r.font.bold = src.bold
    if len(paras) > 34:
        _set_paragraph_text_preserve_format(paras[34], phone)

    # Add a date line in the blank paragraph just under the title (P12),
    # without inserting new paragraphs or disturbing the letter layout.
    if len(paras) > 12 and not paras[12].text.strip():
        _set_paragraph_text_preserve_format(paras[12], f"Date: {doc_date}")

    doc.save(str(out_path))


def fill_methods(template_path: Path, out_path: Path, *,
                 signatory: str, doc_date_month_year: str) -> None:
    """
    Copy + fill the PM16 Bolt Tightening Methods.

    PM16 is a static method statement with no body placeholders. The only
    personalisation is the footer (template: 'Riley Brooks October 2023').
    We update that footer to the current signatory + month/year, staying close
    to the source document.
    """
    shutil.copy2(template_path, out_path)
    doc = Document(str(out_path))
    footer_text = f"{signatory} {doc_date_month_year}"
    for section in doc.sections:
        for fp in section.footer.paragraphs:
            if fp.text.strip():
                _set_paragraph_text_preserve_format(fp, footer_text)
                break
    doc.save(str(out_path))


# --- PDF rendering -----------------------------------------------------------

def render_pdf(docx_path: Path) -> Path:
    """Render a DOCX to PDF alongside it using docx2pdf (Word COM)."""
    from docx2pdf import convert
    pdf_path = docx_path.with_suffix(".pdf")
    convert(str(docx_path), str(pdf_path))
    if not pdf_path.exists():
        raise RuntimeError(f"PDF was not produced: {pdf_path}")
    return pdf_path


# --- Orchestration -----------------------------------------------------------

def generate(*, project_number: str, scope: str, contractor: str,
             project_name: str, signatory: str, role: str, phone: str,
             create_scope: bool, dest_subfolder: str | None,
             make_pdf: bool) -> dict:
    """Run the full generation. Returns a dict of result paths."""
    project_folder = find_project_folder(project_number)
    scope_folder = resolve_scope_folder(project_folder, scope, create_scope)

    dest = scope_folder
    if dest_subfolder:
        dest = scope_folder / dest_subfolder
        dest.mkdir(parents=True, exist_ok=True)

    if not PM17_TEMPLATE.exists():
        raise FileNotFoundError(f"PM17 template not found: {PM17_TEMPLATE}")
    if not PM16_TEMPLATE.exists():
        raise FileNotFoundError(f"PM16 template not found: {PM16_TEMPLATE}")

    today = _dt.date.today()
    doc_date = today.strftime("%d %B %Y")
    month_year = today.strftime("%B %Y")

    cert_docx = dest / f"{project_number} - Bolt tightening Certificate - {scope}.docx"
    methods_docx = dest / f"{project_number} - Bolt Tightening Methods - {scope}.docx"

    fill_certificate(
        PM17_TEMPLATE, cert_docx,
        contractor=contractor, project_name=project_name, scope=scope,
        signatory=signatory, role=role, phone=phone, doc_date=doc_date,
    )
    fill_methods(
        PM16_TEMPLATE, methods_docx,
        signatory=signatory, doc_date_month_year=month_year,
    )

    result = {
        "project_folder": str(project_folder),
        "scope_folder": str(scope_folder),
        "dest": str(dest),
        "cert_docx": str(cert_docx),
        "methods_docx": str(methods_docx),
        "cert_pdf": None,
        "methods_pdf": None,
        "doc_date": doc_date,
    }

    if make_pdf:
        result["cert_pdf"] = str(render_pdf(cert_docx))
        result["methods_pdf"] = str(render_pdf(methods_docx))

    return result


# --- CLI ---------------------------------------------------------------------

def main(argv=None) -> int:
    ap = argparse.ArgumentParser(description="Generate Dunsteel bolt cert + methods.")
    ap.add_argument("--project", required=True, help="Project number, e.g. 501")
    ap.add_argument("--scope", help="QA scope, e.g. 'Level 2', 'Stair 1', 'HV Room'")
    ap.add_argument("--list-scopes", action="store_true",
                    help="List available QA scope folders and exit")
    ap.add_argument("--contractor", default="Northbridge Constructions",
                    help="Head contractor for the Attn line")
    ap.add_argument("--project-name", default="Riverside Eastern",
                    help="Project name for the regarding line")
    ap.add_argument("--signatory", default=DEFAULT_SIGNATORY)
    ap.add_argument("--role", default=DEFAULT_ROLE)
    ap.add_argument("--phone", default=DEFAULT_PHONE)
    ap.add_argument("--create-scope", action="store_true",
                    help="Create the scope folder if missing (skill passes this "
                         "only after user confirmation)")
    ap.add_argument("--dest-subfolder", default=None,
                    help="Write into [scope]/<sub> instead of [scope] (test use)")
    ap.add_argument("--no-pdf", action="store_true", help="Skip PDF rendering")
    args = ap.parse_args(argv)

    if args.list_scopes:
        folder = find_project_folder(args.project)
        scopes = list_scope_folders(folder)
        print(f"Project: {folder.name}")
        print("Available QA scopes:")
        for s in scopes:
            print(f"  - {s}")
        return 0

    if not args.scope:
        ap.error("--scope is required unless --list-scopes is given")

    result = generate(
        project_number=args.project,
        scope=args.scope,
        contractor=args.contractor,
        project_name=args.project_name,
        signatory=args.signatory,
        role=args.role,
        phone=args.phone,
        create_scope=args.create_scope,
        dest_subfolder=args.dest_subfolder,
        make_pdf=not args.no_pdf,
    )

    print("Generated:")
    print(f"  Project folder : {result['project_folder']}")
    print(f"  Scope folder   : {result['scope_folder']}")
    print(f"  Output dir     : {result['dest']}")
    print(f"  Cert DOCX      : {result['cert_docx']}")
    print(f"  Methods DOCX   : {result['methods_docx']}")
    if result["cert_pdf"]:
        print(f"  Cert PDF       : {result['cert_pdf']}")
    if result["methods_pdf"]:
        print(f"  Methods PDF    : {result['methods_pdf']}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
